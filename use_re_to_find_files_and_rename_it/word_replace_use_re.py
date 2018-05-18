# coding:utf-8
# !/usr/env/bin python

'''
读取word文件，使用正则法则将其中的文字替换。
替换Word中的表格信息的文字。
'''

# 读取docx中的文本代码示例
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH


class Elec_interface_Strut:
    id = 0
    sn = ""
    product_description = ""
    fun_des = ""
    remarks = ""


def replace_something(filename="12.docx"):
    # 获取文档对象
    doc_file = docx.Document(filename)
    print(u"段落数:" + str(len(doc_file.paragraphs)))  # 段落数为13，每个回车隔离一段

    # 输出每一段的内容
    for para in doc_file.paragraphs:
        print(para.text)

    #
    # # 输出段落编号及段落内容
    # for i in range(len(file.paragraphs)):
    #     print(u"第" + str(i) + u"段的内容是：" + file.paragraphs[i].text)
    print dir(doc_file)
    t = doc_file.tables[1]
    # print t.rows[0].cells[1].text
    for each in doc_file.tables:
        if each.rows[0].cells[0].text == u"接插件编号：":
            if each.rows[2].cells[1].text != "":
                print each.rows[0].cells[1].text, " " * 6, each.rows[2].cells[1].text
                if each.rows[2].cells[1].text == u"D-Sub9，Male":
                    each.rows[2].cells[1].text = u"D-Sub9M"
                    p = each.rows[2].cells[1].paragraphs[0]
                    # 令单元格中的内容居中
                    p_format = p.paragraph_format
                    p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # print "*"*100

                    pass
            else:
                print each.rows[0].cells[1].text, " " * 6, each.rows[1].cells[1].text
    doc_file.save("12.docx")


def get_table_info(filename=ur"\\172.16.78.59\f\sharefiles\KrF电气互联总图素材\20180511 KrF-IL-ICD 照明系统接口定义.docx"):
    # 获取文档对象
    doc_file = docx.Document(filename)

    Elec_interface_info = Elec_interface_Strut()
    Elec_interface_info.id = 0
    for each in doc_file.tables:
        if each.rows[0].cells[0].text == u"接插件编号：":
            Elec_interface_info.id += 1
            Elec_interface_info.sn = each.rows[0].cells[1].text
            Elec_interface_info.product_description = each.rows[2].cells[1].text
            Elec_interface_info.fun_des = each.rows[4].cells[1].text
            Elec_interface_info.remarks = u"无"

            print u"{0}\t{1}\t{2}\t{3}\t{4}".format(Elec_interface_info.id,Elec_interface_info.sn,Elec_interface_info.product_description,Elec_interface_info.fun_des,Elec_interface_info.remarks)
            # print u"{0}".format( Elec_interface_info.sn)



get_table_info()
